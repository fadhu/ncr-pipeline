"""
Multi-Agent Pipeline: Local NCR Files → Structured CAPA Table
==============================================================
Reads Non-Conformance Report files (PDF, DOCX, scanned images) from a local
folder, processes them through specialized AI agents, and outputs a structured
Excel table with categorized Issue, Corrective Action, and Preventive Action tags.

Supports:
  - PDFs (text-based and scanned)
  - Word documents (.docx)
  - Scanned images (.jpg, .jpeg, .png, .tiff, .bmp)
  - Mixed structured forms and free-form reports

Requirements:
    pip install anthropic pydantic openpyxl tenacity python-docx pdfplumber \
                pytesseract pdf2image Pillow

    System dependencies:
        sudo apt install tesseract-ocr poppler-utils

Usage:
    python ncr_pipeline.py --input-folder ./ncr_files --output ncr_report.xlsx
    python ncr_pipeline.py --input-folder ./ncr_files --output ncr_report.xlsx --concurrency 3

Folder structure example:
    ncr_files/
    ├── NCR-2024-0451.pdf          (text PDF)
    ├── NCR-2024-0452.docx         (Word doc)
    ├── NCR-2024-0453_scan.pdf     (scanned PDF)
    ├── NCR-2024-0454.jpg          (photo of form)
    └── NCR-2024-0455.png          (screenshot)
"""

from __future__ import annotations

import os
import re
import json
import asyncio
import logging
import base64
from pathlib import Path
from datetime import datetime
from typing import Optional
from enum import Enum

from dotenv import load_dotenv
load_dotenv()

from pydantic import BaseModel, Field
from tenacity import retry, stop_after_attempt, wait_exponential
from google import genai
from google.genai import types as genai_types

# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Configuration
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

GEMINI_MODEL = "gemini-2.5-flash"

# Supported file extensions
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp", ".webp"}
ALL_SUPPORTED = PDF_EXTENSIONS | DOCX_EXTENSIONS | IMAGE_EXTENSIONS

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s │ %(name)-18s │ %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("ncr-pipeline")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Taxonomy Enums — customise to your organisation
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class IssueCategory(str, Enum):
    MATERIAL_DEFECT = "Material Defect"
    WORKMANSHIP_ERROR = "Workmanship Error"
    DESIGN_DEVIATION = "Design Deviation"
    SPEC_NON_COMPLIANCE = "Specification Non-Compliance"
    SAFETY_VIOLATION = "Safety Violation"
    ENVIRONMENTAL_NC = "Environmental Non-Compliance"
    DOCUMENTATION_GAP = "Documentation Gap"
    PROCESS_DEVIATION = "Process Deviation"
    OTHER = "Other"


class CorrectiveActionCategory(str, Enum):
    REWORK_REPAIR = "Rework / Repair"
    REPLACEMENT = "Replacement"
    REINSPECTION = "Reinspection"
    USE_AS_IS = "Use-As-Is (with concession)"
    REJECT_SCRAP = "Reject / Scrap"
    DESIGN_REVISION = "Design Revision"
    IMMEDIATE_CONTAINMENT = "Immediate Containment"
    OTHER = "Other"


class PreventiveActionCategory(str, Enum):
    PROCESS_UPDATE = "Process Update / SOP Revision"
    TRAINING = "Training / Competency"
    SUPPLIER_ACTION = "Supplier Corrective Action"
    DESIGN_CHANGE = "Design Change"
    INSPECTION_ENHANCEMENT = "Inspection Enhancement"
    ROOT_CAUSE_ELIMINATION = "Root Cause Elimination"
    SYSTEM_UPGRADE = "System / Tool Upgrade"
    OTHER = "Other"


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Pydantic Schemas
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class IngestedFile(BaseModel):
    """Raw file info before text extraction."""
    file_path: str
    file_name: str
    file_type: str  # "pdf", "docx", "image", "scanned_pdf"
    raw_text: str = ""
    page_count: int = 0
    has_images: bool = False


class ParsedNCR(BaseModel):
    """Normalised text extracted from the NCR by the parser agent."""
    source_file: str
    ncr_id: str
    project_name: str
    date_raised: str
    date_resolved: str = ""
    days_open: Optional[int] = None
    location: Optional[str] = None
    discipline: Optional[str] = None
    raised_by: Optional[str] = None
    contractor: Optional[str] = None
    issue_text: str
    response_text: str
    closeout_text: str
    combined_context: str
    metadata: dict = Field(default_factory=dict)


class IssueClassification(BaseModel):
    category: IssueCategory
    subcategory: Optional[str] = None
    summary: str
    confidence: float = Field(ge=0.0, le=1.0)
    reasoning: str


class CorrectiveActionTag(BaseModel):
    category: CorrectiveActionCategory
    summary: str
    actions_taken: list[str]
    confidence: float = Field(ge=0.0, le=1.0)


class PreventiveActionTag(BaseModel):
    category: PreventiveActionCategory
    summary: str
    actions_recommended: list[str]
    confidence: float = Field(ge=0.0, le=1.0)


class StructuredNCRRow(BaseModel):
    """Final structured output row for the CAPA table."""
    ncr_id: str
    source_file: str
    project_name: str
    date_raised: str
    date_resolved: str
    days_open: Optional[int]
    discipline: Optional[str]
    location: Optional[str]
    raised_by: Optional[str]
    contractor: Optional[str]
    issue_description: str
    issue_category: str
    issue_subcategory: Optional[str]
    corrective_action_summary: str
    corrective_action_category: str
    corrective_actions_detail: str
    preventive_action_summary: str
    preventive_action_category: str
    preventive_actions_detail: str
    overall_confidence: float
    processing_timestamp: str


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Agent Base Class
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class BaseAgent:
    """Base class for LLM-powered agents."""

    def __init__(self, name: str, client: genai.Client):
        self.name = name
        self.client = client
        self.logger = logging.getLogger(f"agent.{name}")

    @retry(stop=stop_after_attempt(3), wait=wait_exponential(min=2, max=30))
    async def _call_llm(self, system_prompt: str, user_content: list | str) -> str:
        """
        Call Gemini with retry logic.
        user_content can be a string or a list of content parts
        (for multi-modal — text + images).
        """
        self.logger.info("Calling LLM...")
        response = await self.client.aio.models.generate_content(
            model=GEMINI_MODEL,
            contents=user_content,
            config=genai_types.GenerateContentConfig(
                system_instruction=system_prompt,
                max_output_tokens=4000,
            ),
        )
        return response.text


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Agent 1: File Ingestion Agent
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class FileIngestionAgent:
    """
    Reads files from a local folder and extracts raw text.
    Handles PDFs (text + scanned), Word docs, and images via OCR.
    """

    def __init__(self):
        self.logger = logging.getLogger("agent.ingestion")

    def discover_files(self, folder_path: str) -> list[Path]:
        """Scan folder (+ one level of subdirectories) for supported NCR files."""
        folder = Path(folder_path)
        if not folder.exists():
            raise FileNotFoundError(f"Folder not found: {folder_path}")

        files = []
        for f in sorted(folder.iterdir()):
            if f.is_file() and f.suffix.lower() in ALL_SUPPORTED:
                files.append(f)
            elif f.is_dir():
                for sub in sorted(f.iterdir()):
                    if sub.is_file() and sub.suffix.lower() in ALL_SUPPORTED:
                        files.append(sub)

        self.logger.info(f"Discovered {len(files)} NCR files in {folder_path}")
        return files

    async def ingest(self, file_path: Path) -> IngestedFile:
        """Extract text from a single file."""
        ext = file_path.suffix.lower()
        self.logger.info(f"Ingesting: {file_path.name} ({ext})")

        if ext in PDF_EXTENSIONS:
            return await self._ingest_pdf(file_path)
        elif ext in DOCX_EXTENSIONS:
            return await self._ingest_docx(file_path)
        elif ext in IMAGE_EXTENSIONS:
            return await self._ingest_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    async def _ingest_pdf(self, file_path: Path) -> IngestedFile:
        """Extract text from PDF — falls back to OCR for scanned pages."""
        import pdfplumber

        text_parts = []
        page_count = 0
        is_scanned = False

        with pdfplumber.open(str(file_path)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                tables = page.extract_tables()
                table_text = ""
                for table in tables:
                    if table:
                        for row in table:
                            table_text += " | ".join(
                                str(cell) if cell else "" for cell in row
                            ) + "\n"
                combined = (page_text + "\n" + table_text).strip()
                text_parts.append(combined)

        full_text = "\n\n--- PAGE BREAK ---\n\n".join(text_parts)

        # If very little text extracted, likely scanned — fall back to OCR
        if len(full_text.strip()) < 100:
            self.logger.info(f"  Low text yield — running OCR on {file_path.name}")
            is_scanned = True
            full_text = await self._ocr_pdf(file_path)

        return IngestedFile(
            file_path=str(file_path),
            file_name=file_path.name,
            file_type="scanned_pdf" if is_scanned else "pdf",
            raw_text=full_text,
            page_count=page_count,
            has_images=is_scanned,
        )

    async def _ocr_pdf(self, file_path: Path) -> str:
        """OCR a scanned PDF using pdf2image + tesseract."""
        from pdf2image import convert_from_path
        import pytesseract

        images = convert_from_path(str(file_path), dpi=300)
        text_parts = []
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img, lang="eng")
            text_parts.append(f"[Page {i+1}]\n{page_text}")
        return "\n\n".join(text_parts)

    async def _ingest_docx(self, file_path: Path) -> IngestedFile:
        """Extract text from Word documents including tables."""
        import docx

        doc = docx.Document(str(file_path))
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    text_parts.append(row_text)

        return IngestedFile(
            file_path=str(file_path),
            file_name=file_path.name,
            file_type="docx",
            raw_text="\n".join(text_parts),
            page_count=1,
        )

    async def _ingest_image(self, file_path: Path) -> IngestedFile:
        """OCR a scanned image (photo of NCR form)."""
        import pytesseract
        from PIL import Image

        img = Image.open(str(file_path))
        if img.mode != "L":
            img = img.convert("L")
        text = pytesseract.image_to_string(img, lang="eng")

        return IngestedFile(
            file_path=str(file_path),
            file_name=file_path.name,
            file_type="image",
            raw_text=text,
            page_count=1,
            has_images=True,
        )


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Agent 2: Document Parser Agent (LLM-powered)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class DocumentParserAgent(BaseAgent):
    """
    Uses Claude to intelligently parse and normalise raw NCR text.
    Handles both structured forms and free-form reports.
    """

    SYSTEM_PROMPT = """You are a construction quality document parser specialising in
Non-Conformance Reports (NCRs). You receive raw extracted text from NCR files which
may be messy (OCR artifacts, table formatting, mixed layouts).

Your job is to extract and normalise the following fields:

1. ncr_id — the NCR reference number (e.g., NCR-2024-0451, QA/NCR/0012, etc.)
2. project_name — the project name
3. date_raised — date the NCR was opened/raised/initiated (ISO format YYYY-MM-DD if possible)
4. date_resolved — date the NCR was closed/resolved/signed-off (ISO format YYYY-MM-DD if possible).
   Look for labels like: "Close Date", "Date Closed", "Resolution Date", "Closeout Date",
   "Date Resolved", "Completion Date", "Sign-off Date", "Accepted Date".
   If the NCR is still open or no close date is found, return ""
5. location — where the non-conformance occurred (building/zone/area/grid)
6. discipline — engineering discipline (Civil, Structural, MEP, Architectural, etc.)
7. raised_by — person or company who raised the NCR
8. contractor — responsible contractor/subcontractor
9. issue_text — the non-conformance description (what went wrong)
10. response_text — the corrective actions / contractor response
11. closeout_text — close-out comments, preventive actions, lessons learned

RULES:
- If a field is not found, return an empty string ""
- Fix OCR errors and expand common abbreviations (conc → concrete, reinf → reinforcement)
- For structured forms, map form field labels to the correct output fields
- For free-form reports, use your judgement to classify paragraphs into the right sections
- The NCR ID might appear as "NCR No.", "NCR #", "NCR Ref", "Reference", etc.
- If the file name contains an NCR-like ID and no ID is found in text, use the file name
- For dates: normalise to YYYY-MM-DD format. Handle common formats like DD/MM/YYYY,
  MM/DD/YYYY, "12 Jan 2024", "January 12, 2024", etc. If ambiguous, prefer DD/MM/YYYY
  (common in construction industry).

Return ONLY a valid JSON object with exactly these keys:
ncr_id, project_name, date_raised, date_resolved, location, discipline, raised_by,
contractor, issue_text, response_text, closeout_text

Do NOT wrap in markdown code blocks."""

    async def run(self, ingested: IngestedFile) -> ParsedNCR:
        self.logger.info(f"Parsing: {ingested.file_name}")

        # For images, use Claude Vision for better accuracy
        if ingested.has_images and ingested.file_type == "image":
            result_text = await self._parse_with_vision(ingested)
        else:
            prompt = f"""File: {ingested.file_name}
File type: {ingested.file_type}
Pages: {ingested.page_count}

--- RAW EXTRACTED TEXT ---
{ingested.raw_text[:12000]}
--- END ---

Extract and normalise the NCR fields from the above text."""
            result_text = await self._call_llm(self.SYSTEM_PROMPT, prompt)

        # Clean and parse JSON
        result_text = result_text.strip()
        result_text = re.sub(r"^```\w*\n?", "", result_text)
        result_text = re.sub(r"\n?```$", "", result_text)

        data = json.loads(result_text)
        ncr_id = data.get("ncr_id", "") or self._extract_id_from_filename(ingested.file_name)

        # Calculate days open
        date_raised = data.get("date_raised", "")
        date_resolved = data.get("date_resolved", "")
        days_open = self._calculate_days_open(date_raised, date_resolved)

        return ParsedNCR(
            source_file=ingested.file_name,
            ncr_id=ncr_id,
            project_name=data.get("project_name", ""),
            date_raised=date_raised,
            date_resolved=date_resolved,
            days_open=days_open,
            location=data.get("location") or None,
            discipline=data.get("discipline") or None,
            raised_by=data.get("raised_by") or None,
            contractor=data.get("contractor") or None,
            issue_text=data.get("issue_text", ""),
            response_text=data.get("response_text", ""),
            closeout_text=data.get("closeout_text", ""),
            combined_context=ingested.raw_text[:8000],
            metadata={"file_type": ingested.file_type, "page_count": ingested.page_count},
        )

    async def _parse_with_vision(self, ingested: IngestedFile) -> str:
        """Send image directly to Gemini Vision for better parsing of scanned forms."""
        self.logger.info(f"  Using vision parsing for {ingested.file_name}")

        with open(ingested.file_path, "rb") as f:
            image_bytes = f.read()

        ext = Path(ingested.file_path).suffix.lower()
        media_type_map = {
            ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".png": "image/png", ".webp": "image/webp",
            ".tiff": "image/tiff", ".tif": "image/tiff",
            ".bmp": "image/bmp",
        }
        media_type = media_type_map.get(ext, "image/jpeg")

        content = [
            genai_types.Part.from_bytes(data=image_bytes, mime_type=media_type),
            f"""This is a scanned/photographed NCR (Non-Conformance Report) form.
File: {ingested.file_name}

Additionally, here is OCR text extracted from the image (may contain errors):
---
{ingested.raw_text[:4000]}
---

Please extract and normalise all NCR fields from this document.
Use BOTH the image and the OCR text to get the most accurate extraction.
Pay special attention to dates — extract BOTH the date raised/opened AND the
date resolved/closed/signed-off. Normalise all dates to YYYY-MM-DD format.""",
        ]
        return await self._call_llm(self.SYSTEM_PROMPT, content)

    @staticmethod
    def _calculate_days_open(date_raised: str, date_resolved: str) -> Optional[int]:
        """Calculate number of days the NCR was open. Returns None if dates missing."""
        if not date_raised:
            return None
        try:
            opened = datetime.strptime(date_raised, "%Y-%m-%d")
            if date_resolved:
                closed = datetime.strptime(date_resolved, "%Y-%m-%d")
            else:
                closed = datetime.utcnow()  # Still open — count to today
            delta = (closed - opened).days
            return max(delta, 0)
        except (ValueError, TypeError):
            return None

    @staticmethod
    def _extract_id_from_filename(filename: str) -> str:
        patterns = [
            r"(NCR[-_\s]?\d{4}[-_]\d{3,5})",
            r"(NCR[-_\s]?\d{3,6})",
            r"(QA[-_/]NCR[-_/]\d{3,6})",
            r"(\d{4}[-_]NCR[-_]\d{3,5})",
        ]
        for pattern in patterns:
            match = re.search(pattern, filename, re.IGNORECASE)
            if match:
                return match.group(1)
        return Path(filename).stem


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Agent 3: Issue Classifier
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class IssueClassifierAgent(BaseAgent):

    SYSTEM_PROMPT = f"""You are a construction quality expert. Classify the non-conformance
issue into exactly ONE of these categories:

{chr(10).join(f"- {c.value}" for c in IssueCategory)}

Return a JSON object with:
- category: exact category string from the list above
- subcategory: optional finer classification (free text)
- summary: 1-2 sentence summary of the non-conformance issue
- confidence: float 0.0-1.0
- reasoning: brief explanation of why this category

Return ONLY valid JSON, no markdown."""

    async def run(self, parsed: ParsedNCR) -> IssueClassification:
        self.logger.info(f"Classifying issue: {parsed.ncr_id}")
        prompt = f"""NCR: {parsed.ncr_id}
Project: {parsed.project_name}
Discipline: {parsed.discipline or 'Unknown'}
Location: {parsed.location or 'Unknown'}

Issue Description:
{parsed.issue_text}

Additional Context:
{parsed.response_text[:2000]}"""

        result = await self._call_llm(self.SYSTEM_PROMPT, prompt)
        result = re.sub(r"^```\w*\n?|```$", "", result.strip())
        return IssueClassification(**json.loads(result))


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Agent 4: Corrective Action Tagger
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class CorrectiveActionTaggerAgent(BaseAgent):

    SYSTEM_PROMPT = f"""You are a construction CAPA specialist. Extract the CORRECTIVE
ACTIONS — the immediate fixes applied to resolve the non-conformance.

Classify into ONE of:
{chr(10).join(f"- {c.value}" for c in CorrectiveActionCategory)}

If no corrective action is explicitly stated, infer the most likely action
based on the issue type and mark confidence lower (0.3-0.5).

Return a JSON object with:
- category: exact category string
- summary: 1-2 sentence summary
- actions_taken: list of specific actions (strings)
- confidence: float 0.0-1.0

Return ONLY valid JSON, no markdown."""

    async def run(self, parsed: ParsedNCR, issue: IssueClassification) -> CorrectiveActionTag:
        self.logger.info(f"Tagging corrective actions: {parsed.ncr_id}")
        prompt = f"""NCR: {parsed.ncr_id}
Issue Category: {issue.category.value}
Issue: {issue.summary}

Response / Corrective Actions Text:
{parsed.response_text}

Close-out Text:
{parsed.closeout_text}

Full Context:
{parsed.issue_text}"""

        result = await self._call_llm(self.SYSTEM_PROMPT, prompt)
        result = re.sub(r"^```\w*\n?|```$", "", result.strip())
        return CorrectiveActionTag(**json.loads(result))


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Agent 5: Preventive Action Tagger
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class PreventiveActionTaggerAgent(BaseAgent):

    SYSTEM_PROMPT = f"""You are a construction CAPA specialist focused on PREVENTION.
Extract PREVENTIVE ACTIONS — systemic changes to prevent recurrence.

Classify into ONE of:
{chr(10).join(f"- {c.value}" for c in PreventiveActionCategory)}

If no explicit preventive action is documented, RECOMMEND appropriate ones
based on the issue, corrective action, and industry best practices.
Mark confidence lower (0.3-0.5) for recommendations vs documented actions.

Return a JSON object with:
- category: exact category string
- summary: 1-2 sentence summary
- actions_recommended: list of preventive measures (strings)
- confidence: float 0.0-1.0

Return ONLY valid JSON, no markdown."""

    async def run(self, parsed: ParsedNCR, issue: IssueClassification,
                  corrective: CorrectiveActionTag) -> PreventiveActionTag:
        self.logger.info(f"Tagging preventive actions: {parsed.ncr_id}")
        prompt = f"""NCR: {parsed.ncr_id}
Issue: {issue.category.value} — {issue.summary}
Corrective Action: {corrective.category.value} — {corrective.summary}

Close-out / Preventive Measures Text:
{parsed.closeout_text}

Full Context:
{parsed.issue_text}
{parsed.response_text}"""

        result = await self._call_llm(self.SYSTEM_PROMPT, prompt)
        result = re.sub(r"^```\w*\n?|```$", "", result.strip())
        return PreventiveActionTag(**json.loads(result))


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Agent 6: Table Structurer
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class TableStructurerAgent:

    def __init__(self):
        self.logger = logging.getLogger("agent.structurer")

    async def run(self, parsed: ParsedNCR, issue: IssueClassification,
                  corrective: CorrectiveActionTag, preventive: PreventiveActionTag) -> StructuredNCRRow:
        self.logger.info(f"Structuring row: {parsed.ncr_id}")
        avg_confidence = round(
            (issue.confidence + corrective.confidence + preventive.confidence) / 3, 3
        )
        return StructuredNCRRow(
            ncr_id=parsed.ncr_id,
            source_file=parsed.source_file,
            project_name=parsed.project_name,
            date_raised=parsed.date_raised,
            date_resolved=parsed.date_resolved,
            days_open=parsed.days_open,
            discipline=parsed.discipline,
            location=parsed.location,
            raised_by=parsed.raised_by,
            contractor=parsed.contractor,
            issue_description=issue.summary,
            issue_category=issue.category.value,
            issue_subcategory=issue.subcategory,
            corrective_action_summary=corrective.summary,
            corrective_action_category=corrective.category.value,
            corrective_actions_detail=" | ".join(corrective.actions_taken),
            preventive_action_summary=preventive.summary,
            preventive_action_category=preventive.category.value,
            preventive_actions_detail=" | ".join(preventive.actions_recommended),
            overall_confidence=avg_confidence,
            processing_timestamp=datetime.utcnow().isoformat(),
        )


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Excel Exporter
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

def export_to_excel(rows: list[StructuredNCRRow], filepath: str):
    """Export structured rows to a professionally formatted Excel workbook."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "NCR CAPA Report"

    header_fill = PatternFill(start_color="1B2A4A", end_color="1B2A4A", fill_type="solid")
    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=10)
    issue_fill = PatternFill(start_color="FFF1F1", end_color="FFF1F1", fill_type="solid")
    ca_fill = PatternFill(start_color="FFF7ED", end_color="FFF7ED", fill_type="solid")
    pa_fill = PatternFill(start_color="EFF6FF", end_color="EFF6FF", fill_type="solid")
    thin_border = Border(
        left=Side(style="thin", color="D1D5DB"),
        right=Side(style="thin", color="D1D5DB"),
        top=Side(style="thin", color="D1D5DB"),
        bottom=Side(style="thin", color="D1D5DB"),
    )

    headers = [
        "NCR ID", "Source File", "Project", "Date Raised", "Date Resolved",
        "Days Open", "Discipline", "Location", "Raised By", "Contractor",
        "Issue Description", "Issue Category", "Issue Subcategory",
        "Corrective Action", "CA Category", "CA Details",
        "Preventive Action", "PA Category", "PA Details",
        "Confidence", "Processed At",
    ]
    issue_cols = {11, 12, 13}
    ca_cols = {14, 15, 16}
    pa_cols = {17, 18, 19}

    # Highlight for overdue / long-open NCRs
    overdue_fill = PatternFill(start_color="FEF2F2", end_color="FEF2F2", fill_type="solid")
    overdue_font = Font(name="Calibri", bold=True, color="DC2626", size=10)

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True, vertical="center")
        cell.border = thin_border
    ws.row_dimensions[1].height = 30

    for i, row in enumerate(rows, 2):
        values = [
            row.ncr_id, row.source_file, row.project_name, row.date_raised,
            row.date_resolved or "OPEN", row.days_open,
            row.discipline, row.location, row.raised_by, row.contractor,
            row.issue_description, row.issue_category, row.issue_subcategory,
            row.corrective_action_summary, row.corrective_action_category,
            row.corrective_actions_detail,
            row.preventive_action_summary, row.preventive_action_category,
            row.preventive_actions_detail,
            row.overall_confidence, row.processing_timestamp,
        ]
        for col, val in enumerate(values, 1):
            cell = ws.cell(row=i, column=col, value=val)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = thin_border
            if col in issue_cols:
                cell.fill = issue_fill
            elif col in ca_cols:
                cell.fill = ca_fill
            elif col in pa_cols:
                cell.fill = pa_fill

        # Highlight "Days Open" red if > 30 days or still open
        days_cell = ws.cell(row=i, column=6)
        resolved_cell = ws.cell(row=i, column=5)
        if (row.days_open is not None and row.days_open > 30) or not row.date_resolved:
            days_cell.fill = overdue_fill
            days_cell.font = overdue_font
            if not row.date_resolved:
                resolved_cell.fill = overdue_fill
                resolved_cell.font = Font(name="Calibri", bold=True, color="DC2626", size=10)

        ws.cell(row=i, column=20).number_format = "0%"

    widths = [16, 22, 22, 13, 13, 10, 14, 16, 14, 16, 40, 24, 20, 40, 22, 36, 40, 22, 36, 10, 22]
    for col, w in enumerate(widths, 1):
        ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = w

    ws.auto_filter.ref = ws.dimensions
    ws.freeze_panes = "A2"

    # ── Summary Sheet ──
    ws2 = wb.create_sheet("Summary")

    # Issue category breakdown
    ws2.cell(row=1, column=1, value="Issue Category Breakdown").font = Font(bold=True, size=12)
    ws2.cell(row=2, column=1, value="Category").font = Font(bold=True)
    ws2.cell(row=2, column=2, value="Count").font = Font(bold=True)
    ws2.cell(row=2, column=3, value="% of Total").font = Font(bold=True)

    cat_counts = {}
    for row in rows:
        cat_counts[row.issue_category] = cat_counts.get(row.issue_category, 0) + 1
    total = len(rows)
    r = 3
    for cat, count in sorted(cat_counts.items(), key=lambda x: -x[1]):
        ws2.cell(row=r, column=1, value=cat)
        ws2.cell(row=r, column=2, value=count)
        ws2.cell(row=r, column=3, value=count / total if total else 0)
        ws2.cell(row=r, column=3).number_format = "0.0%"
        r += 1

    # Resolution time stats
    r += 2
    ws2.cell(row=r, column=1, value="Resolution Time Analysis").font = Font(bold=True, size=12)
    r += 1

    open_ncrs = [row for row in rows if not row.date_resolved]
    closed_ncrs = [row for row in rows if row.date_resolved]
    days_list = [row.days_open for row in closed_ncrs if row.days_open is not None]

    stats = [
        ("Total NCRs Processed", total),
        ("Open NCRs (unresolved)", len(open_ncrs)),
        ("Closed NCRs", len(closed_ncrs)),
        ("Avg. Days to Resolution", round(sum(days_list) / len(days_list), 1) if days_list else "N/A"),
        ("Min Days to Resolution", min(days_list) if days_list else "N/A"),
        ("Max Days to Resolution", max(days_list) if days_list else "N/A"),
        ("NCRs Open > 30 Days", sum(1 for row in rows if row.days_open and row.days_open > 30)),
    ]
    for label, val in stats:
        ws2.cell(row=r, column=1, value=label)
        ws2.cell(row=r, column=2, value=val)
        if isinstance(val, int) and "Open" in label:
            ws2.cell(row=r, column=2).font = Font(bold=True, color="DC2626")
        r += 1

    # List open NCRs
    if open_ncrs:
        r += 1
        ws2.cell(row=r, column=1, value="Open NCRs Detail").font = Font(bold=True, size=12)
        r += 1
        ws2.cell(row=r, column=1, value="NCR ID").font = Font(bold=True)
        ws2.cell(row=r, column=2, value="Date Raised").font = Font(bold=True)
        ws2.cell(row=r, column=3, value="Days Open").font = Font(bold=True)
        ws2.cell(row=r, column=4, value="Project").font = Font(bold=True)
        ws2.cell(row=r, column=5, value="Issue Category").font = Font(bold=True)
        r += 1
        for ncr in sorted(open_ncrs, key=lambda x: -(x.days_open or 0)):
            ws2.cell(row=r, column=1, value=ncr.ncr_id)
            ws2.cell(row=r, column=2, value=ncr.date_raised)
            ws2.cell(row=r, column=3, value=ncr.days_open)
            if ncr.days_open and ncr.days_open > 30:
                ws2.cell(row=r, column=3).font = Font(bold=True, color="DC2626")
            ws2.cell(row=r, column=4, value=ncr.project_name)
            ws2.cell(row=r, column=5, value=ncr.issue_category)
            r += 1

    ws2.column_dimensions["A"].width = 30
    ws2.column_dimensions["B"].width = 16
    ws2.column_dimensions["C"].width = 14
    ws2.column_dimensions["D"].width = 24
    ws2.column_dimensions["E"].width = 26

    wb.save(filepath)
    logger.info(f"✅ Exported {len(rows)} rows → {filepath}")


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# Pipeline Orchestrator
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

class NCRPipeline:
    """
    Orchestrates the full 6-agent pipeline.

    Usage:
        pipeline = NCRPipeline()
        results = await pipeline.run("./ncr_files")
        pipeline.export(results, "ncr_report.xlsx")
    """

    def __init__(self, gemini_api_key: Optional[str] = None):
        self.gemini_client = genai.Client(
            api_key=gemini_api_key or os.getenv("GEMINI_API_KEY")
        )
        self.ingestor = FileIngestionAgent()
        self.parser = DocumentParserAgent("parser", self.gemini_client)
        self.classifier = IssueClassifierAgent("classifier", self.gemini_client)
        self.ca_tagger = CorrectiveActionTaggerAgent("ca_tagger", self.gemini_client)
        self.pa_tagger = PreventiveActionTaggerAgent("pa_tagger", self.gemini_client)
        self.structurer = TableStructurerAgent()

    async def process_single_file(self, file_path: Path) -> Optional[StructuredNCRRow]:
        """Process one NCR file through all 6 agents."""
        logger.info(f"━━━ Processing: {file_path.name} ━━━")
        try:
            ingested = await self.ingestor.ingest(file_path)
            if not ingested.raw_text.strip():
                logger.warning(f"  ⚠ No text extracted — skipping")
                return None

            parsed = await self.parser.run(ingested)
            logger.info(f"  NCR ID: {parsed.ncr_id}")

            issue = await self.classifier.run(parsed)
            logger.info(f"  Issue: {issue.category.value} ({issue.confidence:.0%})")

            corrective = await self.ca_tagger.run(parsed, issue)
            logger.info(f"  CA:    {corrective.category.value} ({corrective.confidence:.0%})")

            preventive = await self.pa_tagger.run(parsed, issue, corrective)
            logger.info(f"  PA:    {preventive.category.value} ({preventive.confidence:.0%})")

            row = await self.structurer.run(parsed, issue, corrective, preventive)
            logger.info(f"  ✓ Confidence: {row.overall_confidence:.0%}")
            return row

        except Exception as e:
            logger.error(f"  ✗ Failed: {file_path.name} — {e}")
            return None

    async def run(self, input_folder: str, concurrency: int = 3) -> list[StructuredNCRRow]:
        """Run the full pipeline on all files in the folder."""
        files = self.ingestor.discover_files(input_folder)
        if not files:
            logger.warning("No supported files found.")
            return []

        logger.info(f"Processing {len(files)} files (concurrency={concurrency})")
        semaphore = asyncio.Semaphore(concurrency)
        results: list[StructuredNCRRow] = []
        failed: list[str] = []

        async def process_with_limit(fp: Path):
            async with semaphore:
                row = await self.process_single_file(fp)
                if row:
                    results.append(row)
                else:
                    failed.append(fp.name)

        await asyncio.gather(*[process_with_limit(f) for f in files])

        logger.info(f"━━━ Pipeline Complete ━━━")
        logger.info(f"  Succeeded: {len(results)}  |  Failed: {len(failed)}")
        if failed:
            logger.info(f"  Failed files: {failed}")

        return sorted(results, key=lambda r: r.ncr_id)

    def export(self, results: list[StructuredNCRRow], filepath: str = "ncr_report.xlsx"):
        export_to_excel(results, filepath)
        return filepath


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
# CLI
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

if __name__ == "__main__":
    import argparse

    ap = argparse.ArgumentParser(
        description="NCR Multi-Agent Pipeline — Local Files → Structured CAPA Table"
    )
    ap.add_argument("--input-folder", "-i", required=True,
                    help="Path to folder containing NCR files (PDF, DOCX, images)")
    ap.add_argument("--output", "-o", default="ncr_report.xlsx",
                    help="Output Excel file path")
    ap.add_argument("--concurrency", "-c", type=int, default=3,
                    help="Max files processed in parallel")
    args = ap.parse_args()

    pipeline = NCRPipeline()
    results = asyncio.run(pipeline.run(args.input_folder, args.concurrency))

    if results:
        pipeline.export(results, args.output)
        print(f"\n✅ Exported {len(results)} NCRs → {args.output}")
    else:
        print("\n⚠️  No NCRs were successfully processed.")
